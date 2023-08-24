from flask import Flask, render_template, request, send_file
from docx import Document

app = Flask(__name__)

def create_timetable(days, times, classes, doc):
    doc.add_heading('College Class Timetable', level=1)

    table = doc.add_table(rows=len(days) + 1, cols=len(times) + 1)
    table.style = 'Table Grid'

    header_row = table.rows[0]
    header_row.cells[0].text = 'Day / Time'
    for i, time in enumerate(times):
        header_row.cells[i+1].text = time

    for j, day in enumerate(days):
        row = table.rows[j+1].cells
        row[0].text = day
        for i, time in enumerate(times):
            subject = classes.get(day, {}).get(time, "")
            row[i+1].text = subject

def get_timetable_input():
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
    times = input("Enter time slots separated by commas (e.g., '9:00 AM - 10:30 AM, 10:45 AM - 12:15 PM'): ").split(',')
    times = [time.strip() for time in times]

    classes = {}
    for day in days:
        classes[day] = {}
        for time in times:
            subject = input(f"Enter subject name for {day}, {time}: ")
            classes[day][time] = subject

    return days, times, classes

def add_course_info_table(doc, course_data):
    doc.add_heading('Course Information', level=1)

    table = doc.add_table(rows=len(course_data) + 1, cols=5)
    table.style = 'Table Grid'

    header_row = table.rows[0]
    header_row.cells[0].text = 'Course Code'
    header_row.cells[1].text = 'Course Name'
    header_row.cells[2].text = 'No. of Hours'
    header_row.cells[3].text = 'Faculty Name'
    header_row.cells[4].text = 'Department'

    for i, (code, name, hours, faculty, department) in enumerate(course_data, start=1):
        row = table.rows[i]
        row.cells[0].text = code
        row.cells[1].text = name
        row.cells[2].text = hours
        row.cells[3].text = faculty
        row.cells[4].text = department
        
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        doc = Document()

        days, times, classes = get_timetable_input()
        create_timetable(days, times, classes, doc)

        num_courses = int(request.form.get('num_courses'))
        course_data = []

        for i in range(num_courses):
            code = request.form.get(f'code_{i}')
            name = request.form.get(f'name_{i}')
            hours = request.form.get(f'hours_{i}')
            faculty = request.form.get(f'faculty_{i}')
            department = request.form.get(f'department_{i}')
            course_data.append((code, name, hours, faculty, department))

        add_course_info_table(doc, course_data)

        doc_path = 'static/college_timetable_and_courses.docx'
        doc.save(doc_path)
        return render_template('index.html', doc_path=doc_path)

    return render_template('index.html', doc_path=None)

if __name__ == "__main__":
    app.run(debug=True)
